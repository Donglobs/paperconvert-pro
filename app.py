from flask import Flask, render_template, request, jsonify, send_file
from PIL import Image
from io import BytesIO
import base64, re, math
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

try:
    from rapidocr_onnxruntime import RapidOCR
    OCR_ENGINE = RapidOCR()
    OCR_ENGINE_NAME = 'RapidOCR'
except Exception as exc:
    OCR_ENGINE = None
    OCR_ENGINE_NAME = f'Unavailable: {exc}'

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024


def b64img(img, quality=94):
    ok, enc = cv2.imencode('.jpg', img, [cv2.IMWRITE_JPEG_QUALITY, quality])
    if not ok:
        raise RuntimeError('Could not encode image')
    return 'data:image/jpeg;base64,' + base64.b64encode(enc.tobytes()).decode()


def decode_data_url(data):
    if ',' in data:
        data = data.split(',', 1)[1]
    raw = base64.b64decode(data)
    arr = np.frombuffer(raw, np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    if img is None:
        raise ValueError('Invalid image')
    return img


def order_points(pts):
    pts = np.asarray(pts, dtype=np.float32).reshape(4, 2)
    s = pts.sum(axis=1)
    d = np.diff(pts, axis=1).reshape(-1)
    return np.array([
        pts[np.argmin(s)], pts[np.argmin(d)],
        pts[np.argmax(s)], pts[np.argmax(d)]
    ], dtype=np.float32)


def clean_page(img, points):
    h, w = img.shape[:2]
    pts = order_points(points)
    tl, tr, br, bl = pts
    width_a = np.linalg.norm(br - bl)
    width_b = np.linalg.norm(tr - tl)
    height_a = np.linalg.norm(tr - br)
    height_b = np.linalg.norm(tl - bl)
    W = int(np.clip(max(width_a, width_b), 1000, 3400))
    H = int(np.clip(max(height_a, height_b), 1300, 4600))
    # Preserve aspect ratio of the selected quadrilateral.
    if W / H > 1.25:
        H = int(W / 0.72)
    dst = np.array([[0, 0], [W - 1, 0], [W - 1, H - 1], [0, H - 1]], dtype=np.float32)
    M = cv2.getPerspectiveTransform(pts, dst)
    warped = cv2.warpPerspective(img, M, (W, H), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
    gray = cv2.cvtColor(warped, cv2.COLOR_BGR2GRAY)
    if gray.shape[0] < 2200:
        scale = min(1.6, 2200 / gray.shape[0])
        gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
    # Remove broad shadows / paper illumination without destroying character strokes.
    background = cv2.GaussianBlur(gray, (0, 0), 35)
    norm = cv2.divide(gray, background, scale=185)
    norm = cv2.normalize(norm, None, 0, 255, cv2.NORM_MINMAX)
    den = cv2.fastNlMeansDenoising(norm, None, 5, 7, 21)
    sharp = cv2.addWeighted(den, 1.25, cv2.GaussianBlur(den, (0, 0), 1.0), -0.25, 0)
    clahe = cv2.createCLAHE(clipLimit=1.6, tileGridSize=(8, 8)).apply(sharp)
    adaptive = cv2.adaptiveThreshold(clahe, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 41, 9)
    otsu = cv2.threshold(clahe, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
    # A softer binarization keeps small printed strokes.
    soft = cv2.GaussianBlur(clahe, (3, 3), 0)
    soft = cv2.threshold(soft, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
    return {
        'preview': b64img(sharp),
        'gray': b64img(sharp),
        'adaptive': b64img(adaptive),
        'otsu': b64img(otsu),
        'soft': b64img(soft),
        'width': int(sharp.shape[1]),
        'height': int(sharp.shape[0])
    }


def ocr_image(img):
    if OCR_ENGINE is None:
        raise RuntimeError('Server OCR engine is unavailable. Check the deployment logs for rapidocr_onnxruntime.')
    result, _ = OCR_ENGINE(img)
    if result is None:
        return []
    rows = []
    for item in result:
        if len(item) < 3:
            continue
        box, text, score = item[0], str(item[1]), float(item[2])
        text = re.sub(r'\s+', ' ', text).strip()
        if not text:
            continue
        pts = np.asarray(box, dtype=np.float32).reshape(-1, 2)
        x0, y0 = pts.min(axis=0)
        x1, y1 = pts.max(axis=0)
        rows.append({'text': text, 'score': score * 100, 'x': float(x0), 'y': float(y0), 'w': float(x1-x0), 'h': float(y1-y0)})
    return rows


def group_lines(words):
    if not words:
        return []
    words = sorted(words, key=lambda x: (x['y'], x['x']))
    lines = []
    for word in words:
        cy = word['y'] + word['h'] / 2
        assigned = None
        for line in lines:
            if abs(cy - line['cy']) <= max(word['h'], line['h']) * 0.55:
                assigned = line
                break
        if assigned is None:
            lines.append({'cy': cy, 'h': word['h'], 'words': [word]})
        else:
            assigned['words'].append(word)
            assigned['cy'] = (assigned['cy'] + cy) / 2
            assigned['h'] = max(assigned['h'], word['h'])
    lines.sort(key=lambda l: l['cy'])
    out = []
    for line in lines:
        ws = sorted(line['words'], key=lambda x: x['x'])
        text = ' '.join(w['text'] for w in ws)
        conf = sum(w['score'] for w in ws) / len(ws)
        out.append({'text': text, 'confidence': conf, 'x': min(w['x'] for w in ws), 'y': min(w['y'] for w in ws)})
    return out


def plausibility(text):
    chars = re.sub(r'[^A-Za-z0-9\s.,;:()\-/%]', '', text)
    words = re.findall(r"[A-Za-z]{2,}", chars)
    if not words:
        return 0
    reasonable = sum(1 for w in words if re.search(r'[aeiouy]', w.lower()) or w.isupper())
    return 100 * reasonable / len(words)


def choose_best(variants):
    candidates = []
    for label, img in variants:
        words = ocr_image(img)
        lines = group_lines(words)
        text = '\n'.join(line['text'] for line in lines)
        if not words:
            candidates.append({'label': label, 'text': '', 'confidence': 0, 'words': 0, 'lines': []})
            continue
        conf = sum(w['score'] for w in words) / len(words)
        useful = len(re.findall(r'[A-Za-z0-9]{2,}', text))
        score = conf * 0.78 + plausibility(text) * 0.17 + min(useful, 120) / 120 * 5
        candidates.append({'label': label, 'text': text, 'confidence': conf, 'words': len(words), 'lines': lines, 'score': score})
    candidates.sort(key=lambda c: c['score'], reverse=True)
    return candidates[0] if candidates else {'label': '', 'text': '', 'confidence': 0, 'words': 0, 'lines': []}


@app.get('/')
def home():
    return render_template('index.html', engine=OCR_ENGINE_NAME)


@app.get('/api/health')
def health():
    return jsonify({'ok': OCR_ENGINE is not None, 'engine': OCR_ENGINE_NAME})


@app.post('/api/clean')
def api_clean():
    data = request.get_json(force=True)
    try:
        img = decode_data_url(data.get('image', ''))
        pts = data.get('points')
        if not isinstance(pts, list) or len(pts) != 4:
            return jsonify({'error': 'Four corner points are required'}), 400
        result = clean_page(img, pts)
        return jsonify(result)
    except Exception as exc:
        return jsonify({'error': str(exc)}), 400


@app.post('/api/ocr')
def api_ocr():
    data = request.get_json(force=True)
    try:
        variants = []
        for key in ('gray', 'adaptive', 'otsu', 'soft'):
            if data.get(key):
                variants.append((key, decode_data_url(data[key])))
        if not variants:
            return jsonify({'error': 'No OCR image variants supplied'}), 400
        best = choose_best(variants)
        if not best['text'].strip():
            return jsonify({'error': 'The OCR engine could not find readable text.', 'candidates': best}), 422
        return jsonify({
            'engine': 'RapidOCR',
            'text': best['text'],
            'confidence': round(best['confidence'], 1),
            'best_pass': best['label'],
            'words': best['words'],
            'lines': best['lines']
        })
    except Exception as exc:
        return jsonify({'error': str(exc)}), 500


def add_doc_paragraph(doc, line):
    line = line.strip()
    if not line:
        return
    m = re.match(r'^(\d+)[.)]\s+(.*)$', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        p.add_run(m.group(2))
        return
    if len(line) <= 90 and (line.isupper() or line.endswith(':')):
        doc.add_heading(line, level=2)
    else:
        doc.add_paragraph(line)


@app.post('/api/export')
def export():
    data = request.get_json(force=True)
    fmt = data.get('format', 'docx')
    title = (data.get('title') or 'Converted Document').strip()
    pages = data.get('pages', [])
    if fmt == 'docx':
        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Inches(.6); sec.bottom_margin = Inches(.6)
        sec.left_margin = Inches(.7); sec.right_margin = Inches(.7)
        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, page in enumerate(pages):
            for line in (page.get('text') or '').splitlines():
                add_doc_paragraph(doc, line)
            if idx < len(pages) - 1:
                doc.add_page_break()
        bio = BytesIO(); doc.save(bio); bio.seek(0)
        return send_file(bio, as_attachment=True, download_name=re.sub(r'[^\w .-]', '_', title)+'.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    styles = getSampleStyleSheet()
    body = ParagraphStyle('body', parent=styles['BodyText'], fontSize=10.5, leading=14, spaceAfter=7)
    head = ParagraphStyle('head', parent=styles['Heading2'], fontSize=14, leading=17, spaceAfter=7)
    bio = BytesIO(); pdf = SimpleDocTemplate(bio, pagesize=A4, leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42)
    story = [Paragraph(title, styles['Title']), Spacer(1, 10)]
    for idx, page in enumerate(pages):
        for line in (page.get('text') or '').splitlines():
            line = line.strip()
            if not line: continue
            safe = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            style = head if (line.isupper() or line.endswith(':')) and len(line) < 90 else body
            story.append(Paragraph(safe, style))
        if idx < len(pages) - 1: story.append(PageBreak())
    pdf.build(story); bio.seek(0)
    return send_file(bio, as_attachment=True, download_name=re.sub(r'[^\w .-]', '_', title)+'.pdf', mimetype='application/pdf')


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
